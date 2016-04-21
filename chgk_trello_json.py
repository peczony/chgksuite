#!/usr/bin/env python
#! -*- coding: utf-8 -*-
from __future__ import unicode_literals
from __future__ import division
import sys
import os
import codecs
import argparse
import requests
import json
import pdb
from collections import defaultdict

# To obtain a token, please copy & paste 
# to the address bar of your browser 
# the following link:
# https://trello.com/1/connect?key=1d4fe71dd193855686196e7768aa4b05&name=Chgk&response_type=token
# Then copy & paste the obtained token 
# to the "token" field in the trello.json
#
# You're almost done. Now, when you are
# browsing your board, in the address bar
# you see a link like 
# https://trello.com/b/d1ISB5RC/-
# Copy and paste the `d1ISB5RC`-like part
# to the "board_id" field in the trello.json
#
# Now you can run the script.

def process_desc(s):
    return s.replace(r'\`','`')

def getlabels(s):
    return {x['name'] for x in s['labels']}

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--si', action='store_true', 
        help="This flag includes card captions in .4s files. "
        "Useful for editing SI files (rather than CHGK)")
    parser.add_argument('--config', default='trello.json',
        help="Use this if you want to use a file named "
        "differently than the default trello.json")
    parser.add_argument('--labels', action='store_true',
        help="Use this if you also want to have lists based on labels.")
    parser.add_argument('--debug', action='store_true',
        help=argparse.SUPPRESS)

    args = parser.parse_args()
    if args.si:
        from docx import Document
    with open(args.config, 'r') as f:
        trello = json.loads(f.read())
    board_id = trello['board_id']
    params = trello['params']

    req = requests.get("https://trello.com/1/boards/{}".format(board_id),
        data=params)
    if req.status_code != 200:
        print('Error: {}'.format(req.text))
        if args.debug:
            pdb.set_trace()
        sys.exit(1)
    
    _lists = defaultdict(lambda: [])

    json_ = json.loads(req.content.decode('utf8'))
    _names = {}
    for list_ in json_['lists']:
        _names[list_['id']] = list_['name']
    for name in _names:
        _names[name] = _names[name].replace('/','_')
    if args.si:
        _docs = defaultdict(lambda: Document('template.docx'))
    for card in json_['cards']:
        if args.si:
            p = _docs[_names[card['idList']]].add_paragraph()
            p.add_run(
                'Тема {}. '.format(
                    len(_lists[_names[card['idList']]])+1
                    )+card['name']).bold = True
            p = _docs[_names[card['idList']]].add_paragraph()
            p = _docs[_names[card['idList']]].add_paragraph()
            p.add_run(process_desc(card['desc']))
            p = _docs[_names[card['idList']]].add_paragraph()
            p = _docs[_names[card['idList']]].add_paragraph()
        _lists[_names[card['idList']]].append(
            ('Тема {}. '.format(
                len(_lists[_names[card['idList']]]) + 1
                )
                + card['name']+'\n\n' if args.si else '')
            + process_desc(card['desc']
            )
            )
        if args.labels:
            for label in getlabels(card):
                _lists[label].append(
                    (card['name'] if args.si else '')
                    + process_desc(card['desc']))
    if args.si:
        for doc in _docs:
            _docs[doc].save('{}.docx'.format(doc))

    for _list in _lists:
        filename = '{}.4s'.format(_list)
        print('outputting {}'.format(filename))
        with codecs.open(filename,'w','utf8') as f:
           for item in _lists[_list]:
              f.write('\n'+item+'\n')

if __name__ == "__main__":
    main()