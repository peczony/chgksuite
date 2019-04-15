#!usr/bin/env python
# -*- coding: utf-8 -*-
from __future__ import unicode_literals
from __future__ import division

try:
    from xmlrpclib import ServerProxy
except ImportError:
    from xmlrpc.client import ServerProxy
import codecs
import contextlib
import logging
import datetime
import hashlib
import os
import random
import re
import shlex
import shutil
import subprocess
import sys
import time
import tempfile
import traceback
import urllib
import dateparser
import pyperclip

try:
    from Tkinter import Tk, Frame, IntVar, Button, Checkbutton, Entry, Label
    import tkFileDialog as filedialog
except ImportError:
    from tkinter import Tk, Frame, IntVar, Button, Checkbutton, Entry, Label
    from tkinter import filedialog
try:
    basestring
except NameError:
    basestring = str
try:
    input = raw_input
except NameError:
    pass

from docx import Document
from docx.shared import Inches
from PIL import Image
import pyimgur

from chgk_common import (
    get_lastdir,
    set_lastdir,
    on_close,
    DummyLogger,
    log_wrap,
    QUESTION_LABELS,
    check_question,
    retry_wrapper_factory,
    bring_to_front,
)
import typotools
from typotools import remove_excessive_whitespace as rew

args = None
im = None
debug = False
console_mode = False
re_url = re.compile(
    r"""((?:[a-z][\w-]+:(?:/{1,3}|[a-z0-9%])|www\d{0,3}[.]"""
    """|[a-z0-9.\-]+[.‌​][a-z]{2,4}/)(?:[^\s<>]+|(([^\s()<>]+|(([^\s<>]+)))*))+"""
    """(?:(([^\s<>]+|(‌​([^\s<>]+)))*)|[^\s`!\[\]{};:'".,<>?«»“”‘’]))""",
    re.DOTALL,
)
re_perc = re.compile(r"(%[0-9a-fA-F]{2})+")
re_scaps = re.compile(
    r"(^|[\s])([\[\]\(\)«»А-Я \u0301`ЁA-Z]{2,})([\s,!\.;:-\?]|$)"
)
re_em = re.compile(r"_(.+?)_")
re_lowercase = re.compile(r"[а-яё]")
re_uppercase = re.compile(r"[А-ЯЁ]")

REQUIRED_LABELS = set(["question", "answer"])
SOURCEDIR = os.path.dirname(os.path.abspath(__file__))
TARGETDIR = os.getcwd()
IMGUR_CLIENT_ID = "8da1bd97da30ac1"

ENC = "utf8" if sys.platform != "win32" else "cp1251"
CONSOLE_ENC = ENC if sys.platform != "win32" else "cp866"

FIELDS = {
    "zachet": "Зачёт: ",
    "nezachet": "Незачёт: ",
    "comment": "Комментарий: ",
    "source": "Источник: ",
    "author": "Автор: ",
}

WHITEN = {
    "handout": False,
    "zachet": True,
    "nezachet": True,
    "comment": True,
    "source": True,
    "author": False,
}


logger = DummyLogger()
retry_wrapper = None


def make_filename(s, ext, nots=False):
    bn = os.path.splitext(os.path.basename(s))[0]
    if nots:
        return bn + "." + ext
    return "{}_{}.{}".format(
        bn, datetime.datetime.now().strftime("%Y%m%dT%H%M"), ext
    )


@contextlib.contextmanager
def make_temp_directory(**kwargs):
    temp_dir = tempfile.mkdtemp(**kwargs)
    yield temp_dir
    shutil.rmtree(temp_dir)


def proportional_resize(tup):
    if max(tup) > 600:
        return tuple([int(x * 600 / max(tup)) for x in tup])
    if max(tup) < 200:
        return tuple([int(x * 200 / max(tup)) for x in tup])
    return tup


def imgsize(imgfile, dimensions="pixels", emsize=25, dpi=120):
    img = Image.open(imgfile)
    width, height = proportional_resize((img.width, img.height))
    if dimensions == "ems":
        return width / emsize, height / emsize
    if dimensions == "inches":
        return width / dpi, height / dpi
    if dimensions == "pixels":
        return width, height
    return width, height


def parseimg(s, dimensions="pixels"):
    width = -1
    height = -1
    sp = s.split()
    imgfile = sp[-1]
    if not os.path.isabs(imgfile):
        if os.path.isfile(os.path.join(TARGETDIR, imgfile)):
            imgfile = os.path.join(TARGETDIR, imgfile)
        elif os.path.isfile(os.path.join(SOURCEDIR, imgfile)):
            imgfile = os.path.join(SOURCEDIR, imgfile)
        else:
            raise Exception(
                "Image file {} not found in {} and {}".format(
                    imgfile, TARGETDIR, SOURCEDIR
                )
            )

    if len(sp) == 1:
        width, height = imgsize(imgfile, dimensions=dimensions)
        return imgfile.replace("\\", "/"), width, height
    else:
        for spsp in sp[:-1]:
            spspsp = spsp.split("=")
            if spspsp[0] == "w":
                width = spspsp[1]
            if spspsp[0] == "h":
                height = spspsp[1]
        return imgfile.replace("\\", "/"), width, height


def partition(alist, indices):
    return [alist[i:j] for i, j in zip([0] + indices, indices + [None])]


def parse_4s_elem(s):
    def find_next_unescaped(ss, index):
        j = index + 1
        while j < len(ss):
            if ss[j] == "\\" and j + 2 < len(ss):
                j += 2
            if ss[j] == ss[index]:
                return j
            j += 1
        return -1

    for gr in re_url.finditer(s):
        gr0 = gr.group(0)
        s = s.replace(gr0, gr0.replace("_", "\u6565_"))

    # for gr in re_scaps.finditer(s):
    #     gr0 = gr.group(0)
    #     s = s.replace(gr0, '(sc '+gr0.lower()+')')

    grs = sorted(
        [match.group(0) for match in re_perc.finditer(s)],
        key=len,
        reverse=True,
    )
    for gr in grs:
        try:
            s = s.replace(gr, urllib.unquote(gr.encode("utf8")).decode("utf8"))
        except:
            logger.debug(
                "error decoding on line {}: {}\n".format(
                    log_wrap(gr), traceback.format_exc()
                )
            )

    i = 0
    topart = []
    while i < len(s):
        if s[i] == "_" and (i == 0 or s[i - 1] not in {"\\", "\u6565"}):
            logger.debug("found _ at {} of line {}".format(i, s))
            topart.append(i)
            if find_next_unescaped(s, i) != -1:
                topart.append(find_next_unescaped(s, i) + 1)
                i = find_next_unescaped(s, i) + 2
                continue
        if (
            s[i] == "("
            and i + len("(img") < len(s)
            and "".join(s[i : i + len("(img")]) == "(img"
        ):
            topart.append(i)
            if typotools.find_matching_closing_bracket(s, i) is not None:
                topart.append(
                    typotools.find_matching_closing_bracket(s, i) + 1
                )
                i = typotools.find_matching_closing_bracket(s, i)
        # if (s[i] == '(' and i + len('(sc') < len(s) and ''.join(s[i:
        #                     i+len('(sc')])=='(sc'):
        #     debug_print('sc candidate')
        #     topart.append(i)
        #     if not typotools.find_matching_closing_bracket(s, i) is None:
        #         topart.append(
        #             typotools.find_matching_closing_bracket(s, i)+1)
        #         i = typotools.find_matching_closing_bracket(s, i)+2
        i += 1

    topart = sorted(topart)

    parts = [
        ["", "".join(x.replace("\u6565", ""))] for x in partition(s, topart)
    ]

    for part in parts:
        if not part[1]:
            continue
        try:
            if part[1][-1] == "_":
                part[1] = part[1][1:]
                part[0] = "em"
            if not part[1]:
                continue
            if part[1][-1] == "_":
                part[1] = part[1][:-1]
                part[0] = "em"
            if not part[1]:
                continue
            if len(part[1]) > 4 and part[1][:4] == "(img":
                if part[1][-1] != ")":
                    part[1] = part[1] + ")"
                part[1] = typotools.remove_excessive_whitespace(part[1][4:-1])
                part[0] = "img"
                logger.debug("found img at {}".format(part[1]))
            if len(part[1]) > 3 and part[1][:4] == "(sc":
                if part[1][-1] != ")":
                    part[1] = part[1] + ")"
                part[1] = typotools.remove_excessive_whitespace(part[1][3:-1])
                part[0] = "sc"
                logger.debug("found img at {}".format(log_wrap(part[1])))
            part[1] = part[1].replace("\\_", "_")
        except:
            sys.stderr.write(
                "Error on part {}: {}".format(
                    log_wrap(part), traceback.format_exc()
                )
            )

    return parts


def parse_4s(s, randomize=False):
    mapping = {
        "#": "meta",
        "##": "section",
        "###": "heading",
        "###LJ": "ljheading",
        "#EDITOR": "editor",
        "#DATE": "date",
        "?": "question",
        "№": "number",
        "№№": "setcounter",
        "!": "answer",
        "=": "zachet",
        "!=": "nezachet",
        "^": "source",
        "/": "comment",
        "@": "author",
        ">": "handout",
    }

    structure = []

    if s[0] == "\ufeff" and len(s) > 1:
        s = s[1:]

    with codecs.open("raw.debug", "w", "utf8") as debugf:
        debugf.write(log_wrap(s.split("\n")))

    for line in s.split("\n"):
        if rew(line) == "":
            structure.append(["", ""])
        else:
            if line.split()[0] in mapping:
                structure.append(
                    [
                        mapping[line.split()[0]],
                        rew(line[len(line.split()[0]) :]),
                    ]
                )
            else:
                if len(structure) >= 1:
                    structure[len(structure) - 1][1] += "\n" + line

    final_structure = []
    current_question = {}
    counter = 1

    if debug:
        with codecs.open("debug1st.debug", "w", "utf8") as debugf:
            debugf.write(log_wrap(structure))

    for element in structure:

        # find list in element

        sp = element[1].split("\n")
        if len(sp) > 1:
            list_candidate = []

            for line in sp:
                if len(rew(line).split()) > 1 and rew(line).split()[0] == "-":
                    list_candidate.append(rew(rew(line)[1:]))

            sp = [
                spsp for spsp in sp if rew(rew(spsp)[1:]) not in list_candidate
            ]

            if len(sp) == 0 or len(sp) == 1 and sp[0] == "":
                element[1] = list_candidate
            else:
                element[1] = (
                    ["\n".join(sp), list_candidate]
                    if len(list_candidate) > 1
                    else "\n".join(element[1].split("\n"))
                )

        if element[0] in QUESTION_LABELS:
            if element[0] in current_question:

                if isinstance(
                    current_question[element[0]], basestring
                ) and isinstance(element[1], basestring):
                    current_question[element[0]] += "\n" + element[1]

                elif isinstance(
                    current_question[element[0]], list
                ) and isinstance(element[1], basestring):
                    current_question[element[0]][0] += "\n" + element[1]

                elif isinstance(
                    current_question[element[0]], basestring
                ) and isinstance(element[1], list):
                    current_question[element[0]] = [
                        element[1][0] + "\n" + current_question[element[0]],
                        element[1][1],
                    ]

                elif isinstance(
                    current_question[element[0]], list
                ) and isinstance(element[1], list):
                    current_question[element[0]][0] += "\n" + element[1][0]
                    current_question[element[0]][1] += element[1][1]
            else:
                current_question[element[0]] = element[1]

        elif element[0] == "":

            if current_question != {} and set(current_question.keys()) != {
                "setcounter"
            }:

                try:
                    assert all(
                        (True if label in current_question else False)
                        for label in REQUIRED_LABELS
                    )
                except AssertionError:
                    logger.error(
                        "Question {} misses "
                        "some of the required fields "
                        "and will therefore "
                        "be omitted.".format(log_wrap(current_question))
                    )
                    continue
                if "setcounter" in current_question:
                    counter = int(current_question["setcounter"])
                if "number" not in current_question:
                    current_question["number"] = counter
                    counter += 1
                final_structure.append(["Question", current_question])

                current_question = {}

        else:
            final_structure.append([element[0], element[1]])

    if current_question != {}:
        try:
            assert all(
                (True if label in current_question else False)
                for label in REQUIRED_LABELS
            )
            if "setcounter" in current_question:
                counter = int(current_question["setcounter"])
            if "number" not in current_question:
                current_question["number"] = counter
                counter += 1
            final_structure.append(["Question", current_question])
        except AssertionError:
            logger.error(
                "Question {} misses "
                "some of the required fields and will therefore "
                "be omitted.".format(log_wrap(current_question))
            )

    if randomize:
        random.shuffle(final_structure, lambda: 0.3)
        i = 1
        for element in final_structure:
            if element[0] == "Question":
                element[1]["number"] = i
                i += 1

    if debug:
        with codecs.open("debug.debug", "w", "utf8") as debugf:
            debugf.write(log_wrap(final_structure))

    for element in final_structure:
        if element[0] == "Question":
            check_question(element[1], logger=logger)

    return final_structure


def gui_get_filetype():
    ch_spoilers = IntVar()
    ch_answers = IntVar()
    ch_rawtex = IntVar()
    if args.nospoilers:
        ch_spoilers.set(0)
    else:
        ch_spoilers.set(1)
    if args.noanswers:
        ch_answers.set(1)
    else:
        ch_answers.set(0)
    if args.rawtex:
        ch_rawtex.set(1)
    else:
        ch_rawtex.set(0)
    root = Tk()
    root.eval("tk::PlaceWindow . center")
    root.grexit = lambda: on_close(root)
    root.ret = None
    root.protocol("WM_DELETE_WINDOW", root.grexit)
    frame = Frame(root)
    frame.pack()
    bottomframe = Frame(root)
    bottomframe.pack(side="bottom")

    def docxreturn():
        root.ret = "docx", ch_spoilers.get(), ch_answers.get(), ch_rawtex.get()
        root.quit()
        root.destroy()

    def texreturn():
        root.ret = "tex", ch_spoilers.get(), ch_answers.get(), ch_rawtex.get()
        root.quit()
        root.destroy()

    def ljreturn():
        root.ret = "lj", ch_spoilers.get(), ch_answers.get(), ch_rawtex.get()
        root.quit()
        root.destroy()

    def chtoggle():
        if ch_spoilers.get() == 0:
            ch_spoilers.set(1)
        else:
            ch_spoilers.set(0)

    def antoggle():
        if ch_answers.get() == 0:
            ch_answers.set(1)
        else:
            ch_answers.set(0)

    def rttoggle():
        if ch_rawtex.get() == 0:
            ch_rawtex.set(1)
        else:
            ch_rawtex.set(0)

    Button(frame, command=docxreturn, text="docx").pack(side="left")
    Button(frame, command=texreturn, text="tex").pack(side="left")
    Button(frame, command=ljreturn, text="LJ").pack(side="left")
    ch = Checkbutton(
        bottomframe,
        text="Spoilers (docx/lj only)",
        variable=ch_spoilers,
        command=chtoggle,
    )
    ans = Checkbutton(
        bottomframe,
        text="No answers (docx only)",
        variable=ch_answers,
        command=antoggle,
    )
    rawtex = Checkbutton(
        bottomframe,
        text="I want to edit tex",
        variable=ch_rawtex,
        command=rttoggle,
    )
    if ch_spoilers.get() == 1:
        ch.select()
    if ch_answers.get() == 1:
        ans.select()
    rawtex.pack(side="bottom")
    ans.pack(side="bottom")
    ch.pack(side="bottom")
    bring_to_front(root)
    root.mainloop()
    return root.ret


def docx_format(el, para, whiten):
    if isinstance(el, list):

        if len(el) > 1 and isinstance(el[1], list):
            docx_format(el[0], para, whiten)
            licount = 0
            for li in el[1]:
                licount += 1

                p = gui_compose.doc.add_paragraph("{}. ".format(licount))
                docx_format(li, p, whiten)
        else:
            licount = 0
            for li in el:
                licount += 1

                p = gui_compose.doc.add_paragraph("{}. ".format(licount))
                docx_format(li, p, whiten)

    if isinstance(el, basestring):
        logger.debug("parsing element {}:".format(log_wrap(el)))

        while "`" in el:
            if el.index("`") + 1 >= len(el):
                el = el.replace("`", "")
            else:
                if el.index("`") + 2 < len(el) and re.search(
                    r"\s", el[el.index("`") + 2]
                ):
                    el = el[: el.index("`") + 2] + "" + el[el.index("`") + 2 :]
                if el.index("`") + 1 < len(el) and re_lowercase.search(
                    el[el.index("`") + 1]
                ):
                    el = (
                        el[: el.index("`") + 1]
                        + ""
                        + el[el.index("`") + 1]
                        + "\u0301"
                        + el[el.index("`") + 2 :]
                    )
                elif el.index("`") + 1 < len(el) and re_uppercase.search(
                    el[el.index("`") + 1]
                ):
                    el = (
                        el[: el.index("`") + 1]
                        + ""
                        + el[el.index("`") + 1]
                        + "\u0301"
                        + el[el.index("`") + 2 :]
                    )
                el = el[: el.index("`")] + el[el.index("`") + 1 :]
        parsed = parse_4s_elem(el)
        images_exist = False

        for run in parsed:
            if run[0] == "img":
                images_exist = True

        for run in parse_4s_elem(el):
            if run[0] == "":
                r = para.add_run(run[1])
                if whiten and not args.nospoilers:
                    r.style = "Whitened"
                # if images_exist:
                #     para = gui_compose.doc.add_paragraph()

            elif run[0] == "em":
                r = para.add_run(run[1])
                r.italic = True
                if whiten and not args.nospoilers:
                    r.style = "Whitened"
                # if images_exist:
                #     para = gui_compose.doc.add_paragraph()

            elif run[0] == "sc":
                r = para.add_run(run[1])
                r.small_caps = True
                if whiten and not args.nospoilers:
                    r.style = "Whitened"
                # if images_exist:
                #     para = gui_compose.doc.add_paragraph()

            elif run[0] == "img":
                imgfile, width, height = parseimg(run[1], dimensions="inches")
                gui_compose.doc.add_picture(
                    imgfile, width=Inches(width), height=Inches(height)
                )
                para = gui_compose.doc.add_paragraph()


def html_format_question(q):
    yapper = htmlyapper
    if "setcounter" in q:
        gui_compose.counter = int(q["setcounter"])
    res = "<strong>Вопрос {}.</strong> {}".format(
        gui_compose.counter if "number" not in q else q["number"],
        yapper(q["question"])
        + ("\n<lj-spoiler>" if not args.nospoilers else ""),
    )
    if "number" not in q:
        gui_compose.counter += 1
    res += "\n<strong>Ответ: </strong>{}".format(yapper(q["answer"]))
    if "zachet" in q:
        res += "\n<strong>Зачёт: </strong>{}".format(yapper(q["zachet"]))
    if "nezachet" in q:
        res += "\n<strong>Незачёт: </strong>{}".format(yapper(q["nezachet"]))
    if "comment" in q:
        res += "\n<strong>Комментарий: </strong>{}".format(
            yapper(q["comment"])
        )
    if "source" in q:
        res += "\n<strong>Источник{}: </strong>{}".format(
            "и" if isinstance(q["source"], list) else "", yapper(q["source"])
        )
    if "author" in q:
        res += "\n<strong>Автор{}: </strong>{}".format(
            "ы" if isinstance(q["author"], list) else "", yapper(q["author"])
        )
    if not args.nospoilers:
        res += "</lj-spoiler>"
    return res


def htmlrepl(zz):
    zz = zz.replace("&", "&amp;")
    zz = zz.replace("<", "&lt;")
    zz = zz.replace(">", "&gt;")

    # while re_scaps.search(zz):
    #     zz = zz.replace(re_scaps.search(zz).group(1),
    #         '\\tsc{'+re_scaps.search(zz).group(1).lower()+'}')

    while "`" in zz:
        if zz.index("`") + 1 >= len(zz):
            zz = zz.replace("`", "")
        else:
            if zz.index("`") + 2 < len(zz) and re.search(
                r"\s", zz[zz.index("`") + 2]
            ):
                zz = zz[: zz.index("`") + 2] + "" + zz[zz.index("`") + 2 :]
            if zz.index("`") + 1 < len(zz) and re_lowercase.search(
                zz[zz.index("`") + 1]
            ):
                zz = (
                    zz[: zz.index("`") + 1]
                    + ""
                    + zz[zz.index("`") + 1]
                    + "&#x0301;"
                    + zz[zz.index("`") + 2 :]
                )
            elif zz.index("`") + 1 < len(zz) and re_uppercase.search(
                zz[zz.index("`") + 1]
            ):
                zz = (
                    zz[: zz.index("`") + 1]
                    + ""
                    + zz[zz.index("`") + 1]
                    + "&#x0301;"
                    + zz[zz.index("`") + 2 :]
                )
            zz = zz[: zz.index("`")] + zz[zz.index("`") + 1 :]

    return zz


def htmlformat(s):
    res = ""
    for run in parse_4s_elem(s):
        if run[0] == "":
            res += htmlrepl(run[1])
        if run[0] == "em":
            res += "<em>" + htmlrepl(run[1]) + "</em>"
        if run[0] == "img":
            imgfile, w, h = parseimg(run[1])
            if os.path.isfile(imgfile):
                # with open(imgfile, 'rb') as f:
                #     imgdata = f.read()
                # imgfile = 'data:image/{ext};base64,{b64}'.format(
                #     ext=os.path.splitext(imgfile)[-1][1:],
                #     b64=base64.b64encode(imgdata))
                uploaded_image = im.upload_image(imgfile, title=imgfile)
                imgfile = uploaded_image.link

            res += '<img{}{} src="{}"/>'.format(
                "" if w == -1 else " width={}".format(w),
                "" if h == -1 else " height={}".format(h),
                imgfile,
            )
    return res


def htmlyapper(e):
    if isinstance(e, basestring):
        return html_element_layout(e)
    elif isinstance(e, list):
        if not any(isinstance(x, list) for x in e):
            return html_element_layout(e)
        else:
            return "\n".join([html_element_layout(x) for x in e])


def html_element_layout(e):
    res = ""
    if isinstance(e, basestring):
        res = htmlformat(e)
        return res
    if isinstance(e, list):
        res = "\n".join(
            [
                "{}. {}".format(en + 1, html_element_layout(x))
                for en, x in enumerate(e)
            ]
        )
        return res


def md5(s):
    return hashlib.md5(s).hexdigest()


def get_chal(lj, passwd):
    chal = None
    chal = retry_wrapper(lj.getchallenge)["challenge"]
    response = md5(
        chal.encode("utf8") + md5(passwd.encode("utf8")).encode("utf8")
    )
    return (chal, response)


def find_heading(structure):
    h_id = -1
    for e, x in enumerate(structure):
        if x[0] == "ljheading":
            return (e, x)
        elif x[0] == "heading":
            h_id = e
    if h_id >= 0:
        return (h_id, structure[h_id])
    return None


def find_tour(structure):
    for e, x in enumerate(structure):
        if x[0] == "section":
            return (e, x)
    return None


def split_into_tours(structure, general_impression=False):
    result = []
    current = []
    mode = "meta"
    for _, element in enumerate(structure):
        if element[0] != "Question":
            if mode == "meta":
                current.append(element)
            elif element[0] == "section":
                result.append(current)
                current = [element]
                mode = "meta"
            else:
                current.append(element)
        else:
            if mode == "meta":
                current.append(element)
                mode = "questions"
            else:
                current.append(element)
    result.append(current)
    globalheading = find_heading(result[0])[1][1]
    result[0][find_heading(result[0])[0]][1] += ". {}".format(
        find_tour(result[0])[1][1]
    )
    for tour in result[1:]:
        if not find_heading(tour):
            tour.insert(
                0,
                [
                    "ljheading",
                    "{}. {}".format(globalheading, find_tour(tour)[1][1]),
                ],
            )
    if general_impression:
        result.append(
            [
                ["ljheading", "{}. Общие впечатления".format(globalheading)],
                [
                    "meta",
                    "В комментариях к этому посту можно "
                    "поделиться общими впечатлениями от вопросов.",
                ],
            ]
        )
    return result


def lj_process(structure):
    final_structure = [{"header": "", "content": ""}]
    i = 0
    heading = ""
    ljheading = ""
    yapper = htmlyapper
    while i < len(structure) and structure[i][0] != "Question":
        if structure[i][0] == "heading":
            final_structure[0]["content"] += "<center>{}</center>".format(
                yapper(structure[i][1])
            )
            heading = yapper(structure[i][1])
        if structure[i][0] == "ljheading":
            # final_structure[0]['header'] = structure[i][1]
            ljheading = yapper(structure[i][1])
        if structure[i][0] == "date":
            final_structure[0]["content"] += "\n<center>{}</center>".format(
                yapper(structure[i][1])
            )
        if structure[i][0] == "editor":
            final_structure[0]["content"] += "\n<center>{}</center>".format(
                yapper(structure[i][1])
            )
        if structure[i][0] == "meta":
            final_structure[0]["content"] += "\n{}".format(
                yapper(structure[i][1])
            )
        i += 1

    if ljheading != "":
        final_structure[0]["header"] = ljheading
    else:
        final_structure[0]["header"] = heading

    for element in structure[i:]:
        if element[0] == "Question":
            final_structure.append(
                {
                    "header": "Вопрос {}".format(
                        element[1]["number"]
                        if "number" in element[1]
                        else gui_compose.counter
                    ),
                    "content": html_format_question(element[1]),
                }
            )
        if element[0] == "meta":
            final_structure.append(
                {"header": "", "content": yapper(element[1])}
            )

    if not final_structure[0]["content"]:
        final_structure[0]["content"] = "Вопросы в комментариях."
    if debug:
        with codecs.open("lj.debug", "w", "utf8") as f:
            f.write(log_wrap(final_structure))
    lj_post(final_structure)


def lj_post(stru):

    lj = ServerProxy("http://www.livejournal.com/interface/xmlrpc").LJ.XMLRPC

    chal, response = get_chal(lj, args.password)

    now = datetime.datetime.now()
    year = now.strftime("%Y")
    month = now.strftime("%m")
    day = now.strftime("%d")
    hour = now.strftime("%H")
    minute = now.strftime("%M")

    params = {
        "username": args.login,
        "auth_method": "challenge",
        "auth_challenge": chal,
        "auth_response": response,
        "subject": stru[0]["header"],
        "event": stru[0]["content"],
        "year": year,
        "mon": month,
        "day": day,
        "hour": hour,
        "min": minute,
    }

    community = args.community
    if community == "":
        params["security"] = "private"
    elif community.startswith("--group"):
        params["security"] = "usemask"
        params["allowmask"] = community.split("--group")[1] or "2"
    else:
        params["usejournal"] = community
    if community == "--group":
        community = ""

    journal = community if community else args.login

    try:
        post = retry_wrapper(lj.postevent, [params])
        ditemid = post["ditemid"]
        logger.info("Created a post")
        logger.debug(log_wrap(post))
        time.sleep(5)

        for _, x in enumerate(stru[1:], start=1):
            chal, response = get_chal(lj, args.password)
            params = {
                "username": args.login,
                "auth_method": "challenge",
                "auth_challenge": chal,
                "auth_response": response,
                "journal": journal,
                "ditemid": ditemid,
                "parenttalkid": 0,
                "body": x["content"],
                "subject": x["header"],
            }
            comment = retry_wrapper(lj.addcomment, [params])
            logger.info("Added a comment")
            logger.debug(log_wrap(comment))
            time.sleep(random.randint(5, 7))

    except:
        sys.stderr.write(
            "Error issued by LJ API: {}".format(traceback.format_exc())
        )
        sys.exit(1)


def lj_post_getdata():
    root = Tk()
    root.login = None
    root.password = None
    root.community = None
    root.grexit = lambda: on_close(root)
    root.eval("tk::PlaceWindow . center")
    root.protocol("WM_DELETE_WINDOW", root.grexit)
    loginbox = Entry(root)
    pwdbox = Entry(root, show="*")
    communitybox = Entry(root)
    ch_split = IntVar()
    ch_genimp = IntVar()
    if args.splittours:
        ch_split.set(1)
    else:
        ch_split.set(0)
    if args.genimp:
        ch_genimp.set(1)
    else:
        ch_genimp.set(0)

    def sptoggle():
        if ch_split.get() == 0:
            ch_split.set(1)
        else:
            ch_split.set(0)

    def gitoggle():
        if ch_genimp.get() == 0:
            ch_genimp.set(1)
        else:
            ch_genimp.set(0)

    def onpwdentry(evt):
        root.login = loginbox.get()
        root.password = pwdbox.get()
        root.community = communitybox.get()
        root.sp = ch_split.get()
        root.gi = ch_genimp.get()
        root.quit()
        root.destroy()

    def onokclick():
        root.login = loginbox.get()
        root.password = pwdbox.get()
        root.community = communitybox.get()
        root.sp = ch_split.get()
        root.gi = ch_genimp.get()
        root.quit()
        root.destroy()

    Label(root, text="Login").pack(side="top")
    loginbox.pack(side="top")
    Label(root, text="Password").pack(side="top")
    pwdbox.pack(side="top")
    Label(root, text="Community (may be blank)").pack(side="top")
    communitybox.pack(side="top")

    pwdbox.bind("<Return>", onpwdentry)
    loginbox.bind("<Return>", onpwdentry)
    communitybox.bind("<Return>", onpwdentry)

    sp = Checkbutton(
        root, text="Split into tours", variable=ch_split, command=sptoggle
    )
    gi = Checkbutton(
        root,
        text="General impression post",
        variable=ch_genimp,
        command=gitoggle,
    )
    if ch_split.get() == 1:
        sp.select()
    if ch_genimp.get() == 1:
        gi.select()
    sp.pack(side="top")
    gi.pack(side="top")

    Button(root, command=onokclick, text="OK").pack(side="top")
    bring_to_front(root)
    root.mainloop()
    return root.login, root.password, root.community, root.sp, root.gi


def baseyapper(e):
    if isinstance(e, basestring):
        return base_element_layout(e)
    elif isinstance(e, list):
        if not any(isinstance(x, list) for x in e):
            return base_element_layout(e)
        else:
            return "\n".join([base_element_layout(x) for x in e])


def baseformat(s):
    res = ""
    for run in parse_4s_elem(s):
        if run[0] == "":
            res += run[1]
        if run[0] == "em":
            res += run[1]
        if run[0] == "img":
            imgfile, w, h = parseimg(run[1], dimensions="ems")
            if os.path.isfile(imgfile):
                im = pyimgur.Imgur(IMGUR_CLIENT_ID)
                uploaded_image = im.upload_image(imgfile, title=imgfile)
                imgfile = uploaded_image.link
            res += "(pic {})".format(imgfile)
    while res.endswith("\n"):
        res = res[:-1]
    return res


def base_element_layout(e):
    res = ""
    if isinstance(e, basestring):
        res = baseformat(e)
        return res
    if isinstance(e, list):
        res = "\n".join([
            "   {}. {}".format(i + 1, base_element_layout(x))
            for i, x in enumerate(e)
        ])
    return res


BASE_MAPPING = {
    "section": "Тур",
    "heading": "Чемпионат",
    "editor": "Редактор",
    "meta": "Инфо"
}
re_date_sep = re.compile(" [—–-] ")


def wrap_date(s):
    s = s.strip()
    parsed = dateparser.parse(s)
    if isinstance(parsed, datetime.datetime):
        parsed = parsed.date()
    if parsed > datetime.date.today():
        parsed = parsed.replace(year=parsed.year - 1)
    formatted = parsed.strftime("%d-%b-%Y")
    return formatted


def base_format_element(pair):
    if pair[0] == "Question":
        return base_format_question(pair[1])
    if pair[0] in BASE_MAPPING:
        return "{}:\n{}\n\n".format(
            BASE_MAPPING[pair[0]], baseyapper(pair[1])
        )
    elif pair[0] == "date":
        re_search = re_date_sep.search(pair[1])
        if re_search:
            gr0 = re_search.group(0)
            dates = pair[1].split(gr0)
            return "Дата:\n{} - {}\n\n".format(
                wrap_date(dates[0]), wrap_date(dates[-1])
            )
        else:
            return "Дата:\n{}\n\n".format(wrap_date(pair[1]))


def output_base(structure, outfile, args):
    result = []
    for pair in structure:
        res = base_format_element(pair)
        if res:
            result.append(res)
    text = "".join(result)
    with codecs.open(outfile, "w", "utf8") as f:
        f.write(text)
    logger.info("Output: {}".format(outfile))
    if args.clipboard:
        pyperclip.copy(text)


def base_format_question(q):
    if "setcounter" in q:
        gui_compose.counter = int(q["setcounter"])
    res = "Вопрос {}:\n{}\n\n".format(
        gui_compose.counter if "number" not in q else q["number"],
        baseyapper(q["question"])
    )
    if "number" not in q:
        gui_compose.counter += 1
    res += "Ответ:\n{}\n\n".format(baseyapper(q["answer"]))
    if "zachet" in q:
        res += "Зачет:\n{}\n\n".format(baseyapper(q["zachet"]))
    if "nezachet" in q:
        res += "Незачет:\n{}\n\n".format(baseyapper(q["zachet"]))
    if "comment" in q:
        res += "Комментарий:\n{}\n\n".format(baseyapper(q["comment"]))
    if "source" in q:
        res += "Источник:\n{}\n\n".format(baseyapper(q["source"]))
    if "author" in q:
        res += "Автор:\n{}\n\n".format(baseyapper(q["author"]))
    return res


def reddityapper(e):
    if isinstance(e, basestring):
        return reddit_element_layout(e)
    elif isinstance(e, list):
        if not any(isinstance(x, list) for x in e):
            return reddit_element_layout(e)
        else:
            return "\n".join([base_element_layout(x) for x in e])


def redditformat(s):
    res = ""
    for run in parse_4s_elem(s):
        if run[0] == "":
            res += run[1]
        if run[0] == "em":
            res += "_{}_".format(run[1])
        if run[0] == "img":
            imgfile, w, h = parseimg(run[1], dimensions="ems")
            if os.path.isfile(imgfile):
                im = pyimgur.Imgur(IMGUR_CLIENT_ID)
                uploaded_image = im.upload_image(imgfile, title=imgfile)
                imgfile = uploaded_image.link
            res += "[картинка]({})".format(imgfile)
    while res.endswith("\n"):
        res = res[:-1]
    res = res.replace("\n", "  \n")
    return res


def reddit_element_layout(e):
    res = ""
    if isinstance(e, basestring):
        res = redditformat(e)
        return res
    if isinstance(e, list):
        res = "\n".join([
            "{}. {}\n".format(i + 1, reddit_element_layout(x))
            for i, x in enumerate(e)
        ])
    return res


def reddit_format_element(pair):
    if pair[0] == "Question":
        return reddit_format_question(pair[1])


def reddit_format_question(q):
    if "setcounter" in q:
        gui_compose.counter = int(q["setcounter"])
    res = "__Вопрос {}__: {}  \n".format(
        gui_compose.counter if "number" not in q else q["number"],
        reddityapper(q["question"])
    )
    if "number" not in q:
        gui_compose.counter += 1
    res += "__Ответ:__ >!{}!<  \n".format(reddityapper(q["answer"]))
    if "zachet" in q:
        res += "__Зачёт:__ >!{}!<  \n".format(reddityapper(q["zachet"]))
    if "nezachet" in q:
        res += "__Незачёт:__ >!{}!<  \n".format(reddityapper(q["zachet"]))
    if "comment" in q:
        res += "__Комментарий:__ >!{}!<  \n".format(reddityapper(q["comment"]))
    if "source" in q:
        res += "__Источник:__ >!{}!<  \n".format(reddityapper(q["source"]))
    if "author" in q:
        res += "__Автор:__ {}  \n".format(reddityapper(q["author"]))
    return res


def output_reddit(structure, outfile, args):
    result = []
    for pair in structure:
        res = reddit_format_element(pair)
        if res:
            result.append(res)
    text = "\n\n".join(result)
    with codecs.open(outfile, "w", "utf8") as f:
        f.write(text)
    logger.info("Output: {}".format(outfile))



def tex_format_question(q):
    yapper = texyapper
    if "setcounter" in q:
        gui_compose.counter = int(q["setcounter"])
    res = (
        "\n\n\\begin{{minipage}}{{\\textwidth}}\\raggedright\n"
        "\\textbf{{Вопрос {}.}} {} \\newline".format(
            gui_compose.counter if "number" not in q else q["number"],
            yapper(q["question"]),
        )
    )
    if "number" not in q:
        gui_compose.counter += 1
    res += "\n\\textbf{{Ответ: }}{} \\newline".format(yapper(q["answer"]))
    if "zachet" in q:
        res += "\n\\textbf{{Зачёт: }}{} \\newline".format(yapper(q["zachet"]))
    if "nezachet" in q:
        res += "\n\\textbf{{Незачёт: }}{} \\newline".format(
            yapper(q["nezachet"])
        )
    if "comment" in q:
        res += "\n\\textbf{{Комментарий: }}{} \\newline".format(
            yapper(q["comment"])
        )
    if "source" in q:
        res += "\n\\textbf{{Источник{}: }}{} \\newline".format(
            "и" if isinstance(q["source"], list) else "", yapper(q["source"])
        )
    if "author" in q:
        res += "\n\\textbf{{Автор: }}{} \\newline".format(yapper(q["author"]))
    res += "\n\\end{minipage}\n"
    return res


def texrepl(zz):
    zz = re.sub(r"{", r"\{", zz)
    zz = re.sub(r"}", r"\}", zz)
    zz = re.sub(r"\\(?![\}\{])", r"{\\textbackslash}", zz)
    zz = re.sub("%", "\%", zz)
    zz = re.sub(r"\$", "\$", zz)
    zz = re.sub("#", "\#", zz)
    zz = re.sub("&", "\&", zz)
    zz = re.sub("_", r"\_", zz)
    zz = re.sub(r"\^", r"{\\textasciicircum}", zz)
    zz = re.sub(r"\~", r"{\\textasciitilde}", zz)
    zz = re.sub(r'((\"(?=[ \.\,;\:\?!\)\]]))|("(?=\Z)))', "»", zz)
    zz = re.sub(r'(((?<=[ \.\,;\:\?!\(\[)])")|((?<=\A)"))', "«", zz)
    zz = re.sub('"', "''", zz)

    for match in sorted(
        [x for x in re_scaps.finditer(zz)],
        key=lambda x: len(x.group(2)),
        reverse=True,
    ):
        zz = zz.replace(
            match.group(2), "\\tsc{" + match.group(2).lower() + "}"
        )

    # while re_scaps.search(zz):
    #     zz = zz.replace(re_scaps.search(zz).group(2),
    #         '\\tsc{'+re_scaps.search(zz).group(2).lower()+'}')

    torepl = [x.group(0) for x in re.finditer(re_url, zz)]
    for s in range(len(torepl)):
        item = torepl[s]
        while item[-1] in typotools.PUNCTUATION:
            item = item[:-1]
        while (
            item[-1] in typotools.CLOSING_BRACKETS
            and typotools.find_matching_opening_bracket(item, -1) is None
        ):
            item = item[:-1]
        while item[-1] in typotools.PUNCTUATION:
            item = item[:-1]
        torepl[s] = item
    torepl = sorted(set(torepl), key=len, reverse=True)
    hashurls = {}
    for s in torepl:
        hashurls[s] = hashlib.md5(s.encode("utf8")).hexdigest()
    for s in sorted(hashurls, key=len, reverse=True):
        zz = zz.replace(s, hashurls[s])
    hashurls = {v: k for k, v in hashurls.items()}
    for s in sorted(hashurls):
        zz = zz.replace(
            s, "\\url{{{}}}".format(hashurls[s].replace("\\\\", "\\"))
        )

    # debug_print('URLS FOR REPLACING: ' +
    #             pprint.pformat(torepl).decode('unicode_escape'))
    # while len(torepl)>0:
    #     s = torepl[0]
    #     debug_print('STRING BEFORE REPLACEMENT: {}'.format(zz))
    #     zz = zz.replace(s, '\\url{'+s+'}')
    #     debug_print('STRING AFTER REPLACEMENT: {}'.format(zz))
    #     torepl.pop(0)

    zz = zz.replace(" — ", "{\\Hair}—{\\hair}")

    while "`" in zz:
        if zz.index("`") + 1 >= len(zz):
            zz = zz.replace("`", "")
        else:
            if zz.index("`") + 2 < len(zz) and re.search(
                r"\s", zz[zz.index("`") + 2]
            ):
                zz = zz[: zz.index("`") + 2] + "" + zz[zz.index("`") + 2 :]
            if zz.index("`") + 1 < len(zz) and re_lowercase.search(
                zz[zz.index("`") + 1]
            ):
                zz = (
                    zz[: zz.index("`") + 1]
                    + ""
                    + zz[zz.index("`") + 1]
                    + "\u0301"
                    + zz[zz.index("`") + 2 :]
                )
            elif zz.index("`") + 1 < len(zz) and re_uppercase.search(
                zz[zz.index("`") + 1]
            ):
                zz = (
                    zz[: zz.index("`") + 1]
                    + ""
                    + zz[zz.index("`") + 1]
                    + "\u0301"
                    + zz[zz.index("`") + 2 :]
                )
            zz = zz[: zz.index("`")] + zz[zz.index("`") + 1 :]

    return zz


def texformat(s):
    res = ""
    for run in parse_4s_elem(s):
        if run[0] == "":
            res += texrepl(run[1])
        if run[0] == "em":
            res += "\\emph{" + texrepl(run[1]) + "}"
        if run[0] == "img":
            imgfile, w, h = parseimg(run[1], dimensions="ems")
            res += (
                "\\includegraphics"
                + "[width={}{}]".format(
                    "10em" if w == -1 else "{}em".format(w),
                    ", height={}em".format(h) if h != -1 else "",
                )
                + "{"
                + imgfile
                + "}"
            )
    while res.endswith("\n"):
        res = res[:-1]
    res = res.replace("\n", "  \\newline \n")
    return res


def texyapper(e):
    if isinstance(e, basestring):
        return tex_element_layout(e)
    elif isinstance(e, list):
        if not any(isinstance(x, list) for x in e):
            return tex_element_layout(e)
        else:
            return "  \n".join([tex_element_layout(x) for x in e])


def tex_element_layout(e):
    res = ""
    if isinstance(e, basestring):
        res = texformat(e)
        return res
    if isinstance(e, list):
        res = """
\\begin{{compactenum}}
{}
\\end{{compactenum}}
""".format(
            "\n".join(["\\item {}".format(tex_element_layout(x)) for x in e])
        )
    return res


def gui_compose(largs, sourcedir=None):

    global im
    global args
    global console_mode
    args = largs
    global __file__  # to fix stupid __file__
    __file__ = os.path.abspath(__file__)  # handling in python 2

    global debug
    global TARGETDIR
    global SOURCEDIR
    global logger
    global retry_wrapper

    if sourcedir:
        SOURCEDIR = sourcedir

    logger = logging.getLogger("composer")
    logger.setLevel(logging.DEBUG)
    fh = logging.FileHandler("composer.log")
    fh.setLevel(logging.DEBUG)
    ch = logging.StreamHandler()
    if args.debug:
        ch.setLevel(logging.DEBUG)
    else:
        ch.setLevel(logging.INFO)
    formatter = logging.Formatter("%(asctime)s | %(message)s")
    fh.setFormatter(formatter)
    ch.setFormatter(formatter)
    logger.addHandler(fh)
    logger.addHandler(ch)

    retry_wrapper = retry_wrapper_factory(logger)

    root = Tk()
    root.withdraw()

    if args.debug:
        debug = True

    argsdict = vars(args)
    logger.debug(log_wrap(argsdict))

    if args.filename and args.filetype:
        if args.filetype == "lj":
            if args.login and args.password:
                console_mode = True
        else:
            console_mode = True

    ld = get_lastdir()
    if not args.filename:
        print("Choose .4s file to load:")
        args.filename = filedialog.askopenfilenames(
            filetypes=[("chgksuite markup files", "*.4s")], initialdir=ld
        )
        if isinstance(args.filename, tuple):
            args.filename = list(args.filename)
    if args.filename:
        if isinstance(args.filename, list):
            ld = os.path.dirname(os.path.abspath(args.filename[0]))
        else:
            ld = os.path.dirname(os.path.abspath(args.filename))
    set_lastdir(ld)
    if not args.filename:
        print("No file specified.")
        sys.exit(1)

    if isinstance(args.filename, list):
        if not args.merge:
            for fn in args.filename:
                TARGETDIR = os.path.dirname(os.path.abspath(fn))
                filename = os.path.basename(os.path.abspath(fn))
                process_file_wrapper(filename)
        else:
            TARGETDIR = os.path.dirname(os.path.abspath(args.filename[0]))
            process_file_wrapper(args.filename)
    else:
        TARGETDIR = os.path.dirname(os.path.abspath(args.filename))
        filename = os.path.basename(os.path.abspath(args.filename))
        process_file_wrapper(filename)


def process_file_wrapper(filename):
    with make_temp_directory(dir=SOURCEDIR) as tmp_dir:
        for fn in [
            args.docx_template,
            os.path.join(SOURCEDIR, "fix-unnumbered-sections.sty"),
            args.tex_header,
        ]:
            shutil.copy(fn, tmp_dir)
        process_file(filename, tmp_dir)
        os.chdir(SOURCEDIR)


def parse_filepath(filepath):
    with codecs.open(filepath, "r", "utf8") as input_file:
        input_text = input_file.read()
    input_text = input_text.replace("\r", "")
    return parse_4s(input_text, randomize=args.randomize)


def make_merged_filename(filelist):
    filelist = [os.path.splitext(os.path.basename(x))[0] for x in filelist]
    prefix = os.path.commonprefix(filelist)
    suffix = "_".join(x[len(prefix) :] for x in filelist)
    return prefix + suffix


def process_file(filename, srcdir):
    global im
    global args
    SOURCEDIR = srcdir
    os.chdir(SOURCEDIR)

    if isinstance(filename, list):
        structure = []
        for x in filename:
            structure.extend(parse_filepath(x))
        filename = make_merged_filename(filename)
    else:
        structure = parse_filepath(os.path.join(TARGETDIR, filename))

    if args.debug:
        with codecs.open(
            make_filename(filename, "dbg", nots=args.nots), "w", "utf8"
        ) as output_file:
            output_file.write(structure)

    if args.filetype is None:
        print("Choose type of export:")
        answer = gui_get_filetype()
        if not answer:
            print("No type of export specified.")
            sys.exit(1)
        args.filetype, spoil, args.noanswers, args.rawtex = answer
        if not args.filetype:
            print("Filetype not specified.")
            sys.exit(1)
        if spoil:
            args.nospoilers = False
        else:
            args.nospoilers = True
        logger.info(
            "Exporting to {}, spoilers are {}...\n".format(
                args.filetype, "off" if args.nospoilers else "on"
            )
        )

    if args.filetype == "docx":

        outfilename = os.path.join(
            SOURCEDIR, make_filename(filename, "docx", nots=args.nots)
        )
        logger.debug(args.docx_template)
        gui_compose.doc = Document(args.docx_template)
        qcount = 0
        logger.debug(log_wrap(structure))

        for element in structure:
            if element[0] == "meta":
                p = gui_compose.doc.add_paragraph()
                docx_format(element[1], p, False)
                gui_compose.doc.add_paragraph()

            if element[0] in ["editor", "date", "heading", "section"]:
                gui_compose.doc.add_paragraph(element[1]).alignment = 1
                gui_compose.doc.add_paragraph()

            if element[0] == "Question":
                q = element[1]
                p = gui_compose.doc.add_paragraph()
                if "number" not in q:
                    qcount += 1
                if "setcounter" in q:
                    qcount = int(q["setcounter"])
                p.add_run(
                    "Вопрос {}. ".format(
                        qcount if "number" not in q else q["number"]
                    )
                ).bold = True

                if "handout" in q:
                    p = gui_compose.doc.add_paragraph()
                    p.add_run("[Раздаточный материал: ")
                    docx_format(q["handout"], p, WHITEN["handout"])
                    p = gui_compose.doc.add_paragraph()
                    p.add_run("]")
                if not args.noparagraph:
                    p = gui_compose.doc.add_paragraph()

                docx_format(q["question"], p, False)
                p = gui_compose.doc.add_paragraph()

                if not args.noanswers:
                    p.add_run("Ответ: ").bold = True
                    docx_format(q["answer"], p, True)

                    for field in [
                        "zachet",
                        "nezachet",
                        "comment",
                        "source",
                        "author",
                    ]:
                        if field in q:
                            p = gui_compose.doc.add_paragraph()
                            if field == "source" and isinstance(
                                q[field], list
                            ):
                                p.add_run("Источники: ").bold = True
                            else:
                                p.add_run(FIELDS[field]).bold = True
                            docx_format(q[field], p, WHITEN[field])

                gui_compose.doc.add_paragraph()

        gui_compose.doc.save(outfilename)
        if os.path.abspath(SOURCEDIR.lower()) != os.path.abspath(
            TARGETDIR.lower()
        ):
            shutil.copy(outfilename, TARGETDIR)
        logger.info(
            "Output: {}".format(
                os.path.join(TARGETDIR, os.path.basename(outfilename))
            )
        )

    if args.filetype == "tex":

        outfilename = os.path.join(
            SOURCEDIR, make_filename(filename, "tex", nots=args.nots)
        )

        gui_compose.counter = 1

        title = ""
        date = ""
        gui_compose.tex = """\\input{@header}
\\begin{document}
""".replace(
            "@header", os.path.basename(args.tex_header)
        )
        firsttour = True
        for element in structure:
            if element[0] == "heading":
                gui_compose.tex += (
                    "\n{{\\huge {}}}\n"
                    "\\vspace{{0.8em}}\n".format(
                        tex_element_layout(element[1])
                    )
                )
            if element[0] == "date":
                gui_compose.tex += (
                    "\n{{\\large {}}}\n"
                    "\\vspace{{0.8em}}\n".format(
                        tex_element_layout(element[1])
                    )
                )
            if element[0] in {"meta", "editor"}:
                gui_compose.tex += "\n{}\n\\vspace{{0.8em}}\n".format(
                    tex_element_layout(element[1])
                )
            elif element[0] == "section":
                gui_compose.tex += "\n{}\\section*{{{}}}\n\n".format(
                    "\\clearpage" if not firsttour else "",
                    tex_element_layout(element[1]),
                )
                firsttour = False
            elif element[0] == "Question":
                gui_compose.tex += tex_format_question(element[1])

        gui_compose.tex += "\\end{document}"

        with codecs.open(outfilename, "w", "utf8") as outfile:
            outfile.write(gui_compose.tex)
        subprocess.call(
            shlex.split(
                'xelatex -synctex=1 -interaction=nonstopmode "{}"'.format(
                    outfilename
                )
            )
        )
        logger.info(
            "Output: {}".format(
                os.path.join(TARGETDIR, os.path.basename(outfilename))
                + "\n"
                + os.path.join(
                    TARGETDIR,
                    os.path.splitext(os.path.basename(outfilename))[0]
                    + ".pdf",
                )
            )
        )
        if os.path.normpath(SOURCEDIR.lower()) != os.path.normpath(
            TARGETDIR.lower()
        ):
            shutil.copy(os.path.splitext(outfilename)[0] + ".pdf", TARGETDIR)
        if args.rawtex and (
            os.path.normpath(SOURCEDIR.lower())
            != os.path.normpath(TARGETDIR.lower())
        ):
            shutil.copy(outfilename, TARGETDIR)
            shutil.copy(args.tex_header, TARGETDIR)
            shutil.copy(
                os.path.join(SOURCEDIR, "fix-unnumbered-sections.sty"),
                TARGETDIR,
            )

    if args.filetype == "lj":

        if not args.community:
            args.community = ""
        if not args.login:
            args.login, args.password, args.community, args.splittours, args.genimp = (
                lj_post_getdata()
            )
            if not args.login:
                print("Login not specified.")
                sys.exit(1)
        elif not args.password:
            import getpass

            args.password = getpass.getpass()

        im = pyimgur.Imgur(IMGUR_CLIENT_ID)

        gui_compose.counter = 1
        if args.splittours:
            tours = split_into_tours(structure, general_impression=args.genimp)
            for tour in tours:
                lj_process(tour)
        else:
            lj_process(structure)

    if args.filetype == "base":
        gui_compose.counter = 1
        outfilename = os.path.join(
            TARGETDIR, make_filename(filename, "txt", nots=args.nots)
        )
        output_base(structure, outfilename, args)

    if args.filetype == "redditmd":
        gui_compose.counter = 1
        outfilename = os.path.join(
            TARGETDIR, make_filename(filename, "md", nots=args.nots)
        )
        output_reddit(structure, outfilename, args)

    if not console_mode:
        input("Press Enter to continue...")


def main():
    print("This program was not designed to run standalone.")
    input("Press Enter to continue...")


if __name__ == "__main__":
    main()
